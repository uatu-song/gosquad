#!/usr/bin/env python3
"""Build a single DOCX manuscript from Book 2 chapter .txt files.

Chapters are renumbered sequentially 1-N in the DOCX only.
Source files are not modified. HTML comments and markdown artifacts
are stripped.

Dabblewriter format:
- Chapter headings: "Chapter N: Title" (Heading 1)
- POV shifts within a chapter: *** (Dabble scene break)
- Same-POV scene transitions: blank line spacing only (no scene break)
"""

import re
from pathlib import Path
from docx import Document
from docx.shared import Pt

BOOK_DIR = Path(__file__).parent

# Source files in manuscript order
SOURCE_FILES = [
    "chapter_01.txt",
    "chapter_02.txt",
    "chapter_03.txt",
    "chapter_04.txt",
    "chapter_05.txt",
    "chapter_06.txt",
    "chapter_07.txt",
    "chapter_08.txt",
    "chapter_09.txt",
    "chapter_10.txt",
    "chapter_11.txt",
    "chapter_12.txt",
    "chapter_13.txt",
    "chapter_14.txt",
    "chapter_15a.txt",
    "chapter_15b.txt",
    "chapter_16.txt",
    "chapter_17.txt",
    "chapter_18.txt",
    "chapter_19.txt",
]

# POV shift locations: source filename → set of line numbers (1-indexed)
# where --- marks an actual perspective change. All other --- are scene
# transitions within the same POV and get blank-line spacing only.
POV_SHIFTS = {
    "chapter_04.txt": {109, 235},   # Bourn → Victor → Ruth
    "chapter_06.txt": {239},         # Tess → Ben
    "chapter_12.txt": {49, 113},     # Ruth → Ben → group
    "chapter_14.txt": {81, 157},     # Leah → Tess → Victor
    "chapter_15b.txt": {261},        # Tess → Korede
    "chapter_19.txt": {41, 55},      # News → Ahdia → Ruth
}


def extract_subtitle(header_line: str) -> str:
    """Extract subtitle from '# Chapter XX: Subtitle' or '# Chapter XXA: Subtitle'."""
    m = re.match(r"#\s*Chapter\s+\d+[A-Da-d]?\s*:\s*(.+)", header_line.strip())
    return m.group(1).strip() if m else ""


def clean_body(text: str) -> str:
    """Strip HTML comments, STRUCT tags, and end-of-chapter markers."""
    text = re.sub(r"<!--.*?-->", "", text, flags=re.DOTALL)
    text = re.sub(r"\*End of Chapter.*?\*", "", text)
    text = re.sub(r"\*End Chapter.*?\*", "", text)
    return text


def process_chapter(filepath: Path) -> tuple[str, list[str]]:
    """Return (subtitle, list_of_paragraph_tokens) from a chapter file.

    Tokens are either plain text, 'SCENE_BREAK' (same-POV transition),
    or 'POV_BREAK' (perspective shift → Dabble scene break).
    """
    raw = filepath.read_text(encoding="utf-8")
    lines = raw.split("\n")
    filename = filepath.name
    pov_lines = POV_SHIFTS.get(filename, set())

    # Extract subtitle from first header line
    subtitle = ""
    body_start = 0
    for i, line in enumerate(lines):
        stripped = line.strip()
        if stripped.startswith("# Chapter"):
            subtitle = extract_subtitle(stripped)
            body_start = i + 1
            break

    # Build paragraph list, tracking original line numbers for POV detection
    paragraphs = []
    for line_idx in range(body_start, len(lines)):
        line = lines[line_idx]
        line_num = line_idx + 1  # 1-indexed
        stripped = line.strip()

        # Skip blanks
        if not stripped:
            continue

        # Scene/POV break detection
        if stripped == "---":
            if line_num in pov_lines:
                paragraphs.append("POV_BREAK")
            else:
                paragraphs.append("SCENE_BREAK")
            continue

        # Skip end-of-chapter / continuation markers
        if stripped.startswith("*End ") or stripped.startswith("**Continues"):
            continue

        # Skip ## Chapter Close headers
        if stripped.startswith("## Chapter Close") or stripped.startswith("## Chapter close"):
            continue

        # Clean HTML comments from inline content
        cleaned = re.sub(r"<!--.*?-->", "", stripped, flags=re.DOTALL).strip()
        if cleaned:
            paragraphs.append(cleaned)

    return subtitle, paragraphs


def add_chapter(doc: Document, chapter_num: int, subtitle: str, paragraphs: list[str]):
    """Add a chapter to the document."""
    # Chapter heading as plain text — Dabble parses "Chapter N: Title"
    if subtitle:
        heading_text = f"Chapter {chapter_num}: {subtitle}"
    else:
        heading_text = f"Chapter {chapter_num}"

    doc.add_paragraph(heading_text)
    doc.add_paragraph()  # blank line after heading

    # Strip leading/trailing breaks
    start = 0
    while start < len(paragraphs) and paragraphs[start] in ("SCENE_BREAK", "POV_BREAK"):
        start += 1
    end = len(paragraphs)
    while end > start and paragraphs[end - 1] in ("SCENE_BREAK", "POV_BREAK"):
        end -= 1

    for para_text in paragraphs[start:end]:
        if para_text == "POV_BREAK":
            # Dabblewriter scene break: *** on its own line
            doc.add_paragraph()
            doc.add_paragraph("***")
            doc.add_paragraph()
        elif para_text == "SCENE_BREAK":
            # Same-POV transition: blank line spacing only
            doc.add_paragraph()
        else:
            # Plain text paragraph — no italic conversion, leave as-is
            doc.add_paragraph(para_text)


def build_manuscript():
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)

    # Process each chapter
    chapter_num = 0
    for filename in SOURCE_FILES:
        filepath = BOOK_DIR / filename
        if not filepath.exists():
            print(f"WARNING: {filename} not found, skipping")
            continue

        chapter_num += 1
        ch_subtitle, paragraphs = process_chapter(filepath)

        # Count break types for reporting
        pov_count = paragraphs.count("POV_BREAK")
        scene_count = paragraphs.count("SCENE_BREAK")
        text_count = len(paragraphs) - pov_count - scene_count

        breaks_info = ""
        if pov_count:
            breaks_info = f", {pov_count} POV break(s)"
        if scene_count:
            breaks_info += f", {scene_count} scene transition(s)"

        print(f"  Ch {chapter_num:2d}: {ch_subtitle} ({filename}, {text_count} paragraphs{breaks_info})")
        add_chapter(doc, chapter_num, ch_subtitle, paragraphs)
        doc.add_paragraph()  # blank separator before next chapter

    output_path = BOOK_DIR / "Book2_Manuscript.docx"
    doc.save(str(output_path))
    print(f"\nSaved: {output_path}")
    print(f"Total chapters: {chapter_num}")


if __name__ == "__main__":
    build_manuscript()
